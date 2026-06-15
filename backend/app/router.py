"""
Router de entrada: validación de seguridad y decisión de ruta.

Para cada fichero subido:
    1. Valida tamaño, tipo MIME real por firma/contenido, no por extensión.
    2. Rechaza ficheros sospechosos o no admitidos.
    3. Extrae texto cuando hay capa textual útil.
    4. Convierte a imágenes cuando el documento solo tiene imagen o el
       texto disponible es de baja calidad.
    5. Devuelve una Entrada lista para el extractor.

Este módulo es la primera línea de defensa. Debe ser estricto: cualquier
formato ambiguo se rechaza en lugar de enviarlo al parser documental.
"""

from __future__ import annotations

import base64
import contextlib
import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
import warnings
import zipfile
from dataclasses import dataclass
from typing import Literal

from .extractor import Entrada
from .parser_sandbox import SandboxExecutionError, ejecutar_en_sandbox, sandbox_activo

logger = logging.getLogger(__name__)


# =============================================================================
# Límites de seguridad
# =============================================================================

TAMANO_MAXIMO_BYTES = 50 * 1024 * 1024      # 50 MB de fichero documental
PAGINAS_MAXIMAS_PDF = 200
PAGINAS_PDF_HIBRIDO = 5
PAGINAS_PDF_VISION_MAX = 20
PDF_HIBRIDO_CON_IMAGENES = os.getenv("PLUMA_PDF_HIBRIDO_CON_IMAGENES", "false").strip().lower() in {"1", "true", "yes", "si", "sí", "on"}
LONGITUD_MAXIMA_TEXTO = int(os.getenv("MAX_LONGITUD_TEXTO_EXTRAIDO", "800000"))

# OCR local. Se ejecuta dentro del contenedor de aplicación, sin servicios
# externos. Tesseract se invoca por CLI para evitar añadir wrappers Python y
# mantener reducida la superficie de dependencias. Si no hay texto suficiente,
# PlumA conserva la ruta visual como fallback.
PLUMA_OCR_LOCAL = os.getenv("PLUMA_OCR_LOCAL", "true").strip().lower() in {"1", "true", "yes", "si", "sí", "on"}
def _ocr_lang_valido(valor: str | None) -> str:
    """Idiomas de Tesseract: códigos de 3 letras unidos por '+' (p. ej.
    'spa+eng'). Es configuración de operador; ante un formato inválido se
    vuelve al valor por defecto seguro en lugar de pasar a tesseract una
    cadena arbitraria."""
    candidato = (valor or "").strip()
    if re.fullmatch(r"[a-z]{3}(\+[a-z]{3})*", candidato):
        return candidato
    logger.warning("PLUMA_OCR_LANG inválido (%r); se usa 'spa+eng'.", valor)
    return "spa+eng"


def _ocr_psm_valido(valor: str | None) -> str:
    """Page segmentation mode de Tesseract: entero 0-13."""
    candidato = (valor or "").strip()
    if re.fullmatch(r"\d{1,2}", candidato) and 0 <= int(candidato) <= 13:
        return candidato
    logger.warning("PLUMA_OCR_PSM inválido (%r); se usa '6'.", valor)
    return "6"


OCR_LANG = _ocr_lang_valido(os.getenv("PLUMA_OCR_LANG", "spa+eng"))
OCR_PSM = _ocr_psm_valido(os.getenv("PLUMA_OCR_PSM", "6"))
OCR_TIMEOUT_SEGUNDOS = max(5, int(os.getenv("PLUMA_OCR_TIMEOUT_SEGUNDOS", "60")))
OCR_MAX_IMAGENES = max(1, int(os.getenv("PLUMA_OCR_MAX_IMAGENES", "20")))
OCR_MIN_CARACTERES = max(20, int(os.getenv("PLUMA_OCR_MIN_CARACTERES", "80")))
OCR_MIN_ALFANUM_RATIO = float(os.getenv("PLUMA_OCR_MIN_ALFANUM_RATIO", "0.35"))

# Imagen de entrada. 40 MP evita bombas razonables sin impedir escaneos
# administrativos grandes. La dimensión máxima reduce cargas patológicas.
DIMENSION_MAXIMA_IMAGEN = 6_000
PIXELS_MAXIMOS_IMAGEN = 40_000_000
BYTES_MAXIMOS_IMAGEN_NORMALIZADA = 12 * 1024 * 1024

# Render PDF. Se calculan píxeles estimados antes de renderizar.
ESCALA_RENDER_PDF = 2.0
PIXELS_MAXIMOS_PAGINA_PDF = 35_000_000
BYTES_MAXIMOS_IMAGEN_PDF = 12 * 1024 * 1024
BYTES_MAXIMOS_TOTAL_IMAGENES = 80 * 1024 * 1024

# DOCX/ZIP. Límites deliberadamente conservadores para evitar zip bombs.
DOCX_MAX_ENTRADAS = 500
DOCX_MAX_DESCOMPRIMIDO = 30 * 1024 * 1024
DOCX_MAX_RATIO_COMPRESION = 100
DOCX_MAX_DOCUMENT_XML = 5 * 1024 * 1024
DOCX_MAX_MEDIA_TOTAL = 20 * 1024 * 1024

TIPOS_ADMITIDOS = {
    "application/pdf":              "pdf",
    "text/plain":                   "texto",
    "text/markdown":                "texto",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "image/jpeg":                   "imagen",
    "image/png":                    "imagen",
    "image/tiff":                   "imagen",
    "image/webp":                   "imagen",
}

Ruta = Literal["texto", "vision", "hibrida"]
PreferenciaEntrada = Literal["auto", "pdf_texto", "pdf_escaneado", "imagen", "texto"]

PREFERENCIAS_ENTRADA_VALIDAS = {"auto", "pdf_texto", "pdf_escaneado", "imagen", "texto"}

def normalizar_preferencia_entrada(valor: str | None) -> PreferenciaEntrada:
    """Normaliza la pista de entrada indicada por la interfaz.

    Esta pista no sustituye a la validación MIME real; solo evita rutas
    costosas cuando el usuario ya sabe si el PDF tiene texto/OCR o si debe
    tratarse como imagen/escaneo.
    """
    candidato = (valor or "auto").strip().lower().replace("-", "_")
    alias = {
        "automatico": "auto",
        "automático": "auto",
        "pdf_ocr": "pdf_texto",
        "pdf_textual": "pdf_texto",
        "pdf_con_ocr": "pdf_texto",
        "pdf_sin_ocr": "pdf_escaneado",
        "pdf_scan": "pdf_escaneado",
        "scan": "pdf_escaneado",
        "escaneo": "pdf_escaneado",
        "image": "imagen",
        "texto_docx": "texto",
    }
    candidato = alias.get(candidato, candidato)
    if candidato not in PREFERENCIAS_ENTRADA_VALIDAS:
        raise ErrorValidacion(f"Preferencia de entrada no admitida: {valor}")
    return candidato  # type: ignore[return-value]



# =============================================================================
# Excepciones
# =============================================================================

class ErrorValidacion(Exception):
    """Se lanza cuando el fichero no supera las comprobaciones de seguridad."""


# =============================================================================
# Resultado
# =============================================================================

@dataclass
class DocumentoProcesado:
    entrada: Entrada
    ruta: Ruta
    nombre_original: str
    tipo_mime: str
    tamano_bytes: int
    paginas: int | None


# =============================================================================
# Validación
# =============================================================================

def validar(contenido: bytes, nombre: str) -> str:
    """
    Valida el fichero y devuelve su tipo MIME real. No se acepta fallback
    por extensión: si la firma/contenido no encaja, se rechaza.
    """
    if len(contenido) == 0:
        raise ErrorValidacion("El fichero está vacío.")
    if len(contenido) > TAMANO_MAXIMO_BYTES:
        raise ErrorValidacion(
            f"El fichero supera el tamaño máximo "
            f"({len(contenido) / 1_048_576:.1f} MB, máximo "
            f"{TAMANO_MAXIMO_BYTES / 1_048_576:.0f} MB)."
        )

    mime = _detectar_mime_real(contenido)

    if mime not in TIPOS_ADMITIDOS:
        raise ErrorValidacion(
            f"Tipo de fichero no admitido: {mime}. "
            f"Formatos soportados: PDF, DOCX, TXT, JPG, PNG, TIFF, WebP."
        )

    if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        _validar_docx_seguro(contenido)

    return mime


def _detectar_mime_real(contenido: bytes) -> str:
    if len(contenido) < 4:
        raise ErrorValidacion("Fichero demasiado pequeño para ser válido.")

    if contenido.startswith(b"%PDF-"):
        return "application/pdf"
    if contenido.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if contenido.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if contenido.startswith((b"II*\x00", b"MM\x00*")):
        return "image/tiff"
    if contenido.startswith(b"RIFF") and len(contenido) >= 12 and contenido[8:12] == b"WEBP":
        return "image/webp"
    if contenido.startswith(b"PK\x03\x04"):
        if _es_docx(contenido):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        raise ErrorValidacion("Fichero ZIP no reconocido como DOCX admitido.")

    if _parece_texto(contenido):
        return "text/plain"

    raise ErrorValidacion(
        "No se ha podido determinar el tipo del fichero. "
        "Asegúrese de que es un PDF, DOCX, TXT o imagen válida."
    )


def _es_docx(contenido: bytes) -> bool:
    try:
        with zipfile.ZipFile(io.BytesIO(contenido)) as z:
            nombres = set(z.namelist())
    except zipfile.BadZipFile:
        return False
    return "[Content_Types].xml" in nombres and "word/document.xml" in nombres


def _validar_docx_seguro(contenido: bytes) -> None:
    """Valida estructura DOCX antes de entregarla a python-docx."""
    try:
        with zipfile.ZipFile(io.BytesIO(contenido)) as z:
            infos = z.infolist()
    except zipfile.BadZipFile:
        raise ErrorValidacion("El documento DOCX está dañado o no es válido.") from None

    if not infos:
        raise ErrorValidacion("El DOCX no contiene entradas internas válidas.")
    if len(infos) > DOCX_MAX_ENTRADAS:
        raise ErrorValidacion(
            f"El DOCX contiene demasiadas entradas internas ({len(infos)}; "
            f"máximo {DOCX_MAX_ENTRADAS})."
        )

    total_descomprimido = 0
    total_media = 0
    tiene_document_xml = False

    for info in infos:
        nombre = info.filename.replace("\\", "/")

        # Rutas internas anómalas. No extraemos el ZIP, pero las rechazamos
        # para evitar ambigüedades y futuras regresiones si se añade extracción.
        if nombre.startswith("/") or ".." in nombre.split("/"):
            raise ErrorValidacion("El DOCX contiene rutas internas no seguras.")

        if nombre.endswith("/"):
            continue

        if "vbaProject" in nombre:
            raise ErrorValidacion(
                "El documento contiene macros. Por seguridad, no se procesan "
                "documentos con macros. Guárdelo como .docx sin macros."
            )

        total_descomprimido += info.file_size
        if total_descomprimido > DOCX_MAX_DESCOMPRIMIDO:
            raise ErrorValidacion(
                "El DOCX se expande demasiado al descomprimirse. "
                "Puede estar dañado o construido de forma maliciosa."
            )

        if info.compress_size > 0:
            ratio = info.file_size / info.compress_size
            if ratio > DOCX_MAX_RATIO_COMPRESION and info.file_size > 1_000_000:
                raise ErrorValidacion(
                    "El DOCX tiene una ratio de compresión anómala. "
                    "Puede tratarse de un ZIP-bomb."
                )

        if nombre == "word/document.xml":
            tiene_document_xml = True
            if info.file_size > DOCX_MAX_DOCUMENT_XML:
                raise ErrorValidacion("El cuerpo XML principal del DOCX es demasiado grande.")

        if nombre.startswith("word/media/"):
            total_media += info.file_size
            if total_media > DOCX_MAX_MEDIA_TOTAL:
                raise ErrorValidacion("El DOCX contiene demasiados recursos multimedia.")

    if not tiene_document_xml:
        raise ErrorValidacion("El DOCX no contiene word/document.xml.")


def _parece_texto(contenido: bytes) -> bool:
    """
    Detección estricta de texto. Evita aceptar binarios aleatorios por el
    simple hecho de que Latin-1 pueda decodificar cualquier byte.
    """
    muestra = contenido[:8192]
    if not muestra:
        return False

    if muestra.startswith((b"\xff\xfe", b"\xfe\xff")):
        # TXT UTF-16 no se soporta en esta versión para mantener simple la
        # puerta de entrada y evitar confusiones con binarios con muchos nulos.
        return False

    for encoding in ("utf-8", "cp1252", "latin-1"):
        try:
            texto = muestra.decode(encoding)
        except UnicodeDecodeError:
            continue

        if "\x00" in texto:
            return False

        controles = sum(1 for c in texto if ord(c) < 32 and c not in "\n\r\t")
        if controles > max(2, len(texto) * 0.01):
            return False

        imprimibles = sum(1 for c in texto if c.isprintable() or c in "\n\r\t")
        if imprimibles / max(1, len(texto)) < 0.95:
            return False

        letras_o_numeros = sum(1 for c in texto if c.isalnum())
        espacios = sum(1 for c in texto if c.isspace())
        if letras_o_numeros < max(8, len(texto) * 0.20):
            return False
        if len(texto) > 200 and espacios == 0:
            return False

        return True

    return False


# =============================================================================
# Evaluación de calidad del OCR
# =============================================================================

def _calidad_ocr(texto: str, num_paginas: int) -> float:
    """Estima la calidad del texto OCR en una escala 0.0-1.0."""
    if not texto or num_paginas == 0:
        return 0.0

    chars_por_pagina = len(texto) / num_paginas
    if chars_por_pagina < 200:
        return 0.0
    densidad = min(1.0, chars_por_pagina / 600)

    alfanum = sum(1 for c in texto if c.isalnum() or c in " .,;:áéíóúñüÁÉÍÓÚÑÜ¿?¡!")
    proporcion_sana = alfanum / max(1, len(texto))

    palabras = [p for p in texto.split() if len(p) > 3 and any(c.isalpha() for c in p)]
    palabras_por_pagina = len(palabras) / num_paginas
    palabras_score = min(1.0, palabras_por_pagina / 80)

    lineas = [linea.strip() for linea in texto.split("\n") if linea.strip()]
    if lineas:
        cortas = sum(1 for linea in lineas if len(linea.split()) <= 2)
        prop_cortas = cortas / len(lineas)
        lineas_score = max(0.0, 1.0 - max(0.0, prop_cortas - 0.4) * 2)
    else:
        lineas_score = 0.0

    return (densidad * 0.3 + proporcion_sana * 0.3
            + palabras_score * 0.25 + lineas_score * 0.15)


# =============================================================================
# Extracción por tipo
# =============================================================================

def _extraer_texto_pdf(contenido: bytes) -> tuple[str, int]:
    """Extrae texto de un PDF. Devuelve (texto, num_paginas)."""
    import pypdf

    try:
        lector = pypdf.PdfReader(io.BytesIO(contenido), strict=False)
    except Exception as e:
        raise ErrorValidacion(f"El PDF no se puede abrir: {e}") from None

    if lector.is_encrypted:
        raise ErrorValidacion("El PDF está cifrado o protegido por contraseña.")

    num_paginas = len(lector.pages)
    if num_paginas <= 0:
        raise ErrorValidacion("El PDF no contiene páginas.")
    if num_paginas > PAGINAS_MAXIMAS_PDF:
        raise ErrorValidacion(
            f"El PDF tiene {num_paginas} páginas; el máximo admitido es "
            f"{PAGINAS_MAXIMAS_PDF}. Divida el documento en piezas más pequeñas."
        )

    partes = []
    for pagina in lector.pages:
        try:
            partes.append(pagina.extract_text() or "")
        except Exception as e:
            logger.warning("Error extrayendo texto de una página: %s", e)
            partes.append("")

    return "\n\n".join(partes).strip(), num_paginas


def _pdf_a_imagenes(contenido: bytes, max_paginas: int = PAGINAS_PDF_VISION_MAX) -> list[bytes]:
    import pypdfium2 as pdfium

    imagenes: list[bytes] = []
    total_bytes = 0

    try:
        pdf = pdfium.PdfDocument(contenido)
    except Exception as e:
        raise ErrorValidacion(f"El PDF no se puede abrir para renderizado: {e}") from None

    try:
        paginas_a_procesar = min(len(pdf), max_paginas)
        for i in range(paginas_a_procesar):
            pagina = pdf[i]
            try:
                ancho_pt, alto_pt = pagina.get_size()
            except Exception:
                ancho_pt, alto_pt = (0, 0)

            if ancho_pt and alto_pt:
                pixeles_estimados = int(ancho_pt * ESCALA_RENDER_PDF) * int(alto_pt * ESCALA_RENDER_PDF)
                if pixeles_estimados > PIXELS_MAXIMOS_PAGINA_PDF:
                    raise ErrorValidacion(
                        f"La página {i + 1} del PDF es demasiado grande para renderizar "
                        f"({pixeles_estimados:,} píxeles estimados)."
                    )

            bitmap = pagina.render(scale=ESCALA_RENDER_PDF).to_pil()
            try:
                if bitmap.width * bitmap.height > PIXELS_MAXIMOS_PAGINA_PDF:
                    raise ErrorValidacion(
                        f"La página {i + 1} del PDF supera el límite de píxeles renderizados."
                    )

                # Convertir a RGB reduce modos exóticos y transparencias complejas.
                if bitmap.mode not in ("RGB", "L"):
                    bitmap = bitmap.convert("RGB")

                buf = io.BytesIO()
                bitmap.save(buf, format="PNG", optimize=True)
                datos = buf.getvalue()
            finally:
                try:
                    bitmap.close()
                except Exception:
                    pass

            if len(datos) > BYTES_MAXIMOS_IMAGEN_PDF:
                raise ErrorValidacion(
                    f"La imagen renderizada de la página {i + 1} es demasiado grande."
                )
            total_bytes += len(datos)
            if total_bytes > BYTES_MAXIMOS_TOTAL_IMAGENES:
                raise ErrorValidacion("El conjunto de imágenes renderizadas del PDF es demasiado grande.")
            imagenes.append(datos)
    finally:
        pdf.close()

    return imagenes


# =============================================================================
# OCR local con Tesseract
# =============================================================================

def _tesseract_disponible() -> bool:
    """Comprueba si el binario de Tesseract está disponible en PATH."""
    return shutil.which("tesseract") is not None


def _texto_ocr_suficiente(texto: str | None) -> bool:
    """Heurística conservadora para decidir si el OCR evita la ruta visual."""
    if not texto:
        return False
    limpio = re.sub(r"\s+", " ", texto).strip()
    if len(limpio) < OCR_MIN_CARACTERES:
        return False
    alfanum = sum(1 for c in limpio if c.isalnum())
    ratio = alfanum / max(1, len(limpio))
    if ratio < OCR_MIN_ALFANUM_RATIO:
        return False
    palabras = [p for p in re.split(r"\s+", limpio) if len(p) >= 3 and any(c.isalpha() for c in p)]
    return len(palabras) >= 8


def _ocr_imagen_tesseract(imagen: bytes, *, indice: int | None = None) -> str:
    """Ejecuta OCR local sobre una imagen ya validada/renderizada."""
    if not PLUMA_OCR_LOCAL:
        return ""
    if not _tesseract_disponible():
        logger.info("OCR local omitido: tesseract no está disponible en PATH")
        return ""

    # Tesseract trabaja mejor con ficheros temporales que con stdin cuando hay
    # formatos variados. No usamos shell y el directorio temporal vive en tmpfs
    # del contenedor.
    tmp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(prefix="pluma_ocr_", suffix=".png", delete=False) as tmp:
            tmp.write(imagen)
            tmp_path = tmp.name

        cmd = [
            "tesseract",
            tmp_path,
            "stdout",
            "-l",
            OCR_LANG,
            "--psm",
            OCR_PSM,
        ]
        proc = subprocess.run(
            cmd,
            check=False,
            capture_output=True,
            timeout=OCR_TIMEOUT_SEGUNDOS,
        )
        if proc.returncode != 0:
            detalle = (proc.stderr or b"").decode("utf-8", errors="replace").strip()
            logger.info(
                "OCR local sin resultado%s: tesseract rc=%s %s",
                f" imagen={indice}" if indice is not None else "",
                proc.returncode,
                detalle[:300],
            )
            return ""
        return (proc.stdout or b"").decode("utf-8", errors="replace").strip()
    except subprocess.TimeoutExpired:
        logger.info("OCR local cancelado por timeout%s", f" imagen={indice}" if indice is not None else "")
        return ""
    except Exception as exc:
        logger.info("OCR local falló%s: %s", f" imagen={indice}" if indice is not None else "", exc)
        return ""
    finally:
        if tmp_path:
            with contextlib.suppress(Exception):
                os.unlink(tmp_path)


def _ocr_imagenes(imagenes: list[bytes], *, nombre_origen: str = "documento") -> str:
    """OCR local sobre una lista de imágenes, manteniendo orden y paginación."""
    if not PLUMA_OCR_LOCAL or not imagenes:
        return ""

    seleccionadas = imagenes[:OCR_MAX_IMAGENES]
    partes: list[str] = []
    for i, img in enumerate(seleccionadas, start=1):
        texto = _ocr_imagen_tesseract(img, indice=i)
        if _texto_ocr_suficiente(texto):
            partes.append(f"[OCR local - {nombre_origen} - página/imagen {i}]\n{texto.strip()}")
        elif texto:
            logger.info(
                "OCR local descartado por baja calidad en %s imagen %d (%d caracteres)",
                nombre_origen,
                i,
                len(texto),
            )

    if len(imagenes) > len(seleccionadas):
        partes.append(
            f"[Aviso técnico] El OCR local procesó {len(seleccionadas)} de {len(imagenes)} "
            "imágenes para mantener límites razonables de tiempo y memoria."
        )

    return "\n\n".join(partes).strip()


def _extraer_texto_docx(contenido: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(contenido))
    except Exception as e:
        raise ErrorValidacion(f"El DOCX no se puede abrir: {e}") from None

    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                partes.append(" | ".join(celdas))

    return "\n".join(partes).strip()


@contextlib.contextmanager
def _max_image_pixels(valor: int):
    """
    Encapsula la asignación a Image.MAX_IMAGE_PIXELS para restaurar el valor
    previo al salir, incluso si hay excepción.

    Image.MAX_IMAGE_PIXELS es global del módulo PIL, no por instancia. En el
    diseño actual el procesamiento ocurre dentro del sandbox (proceso hijo),
    por lo que no había riesgo real, pero encapsularlo elimina cualquier
    posibilidad de fuga si en una versión futura se desactivara el sandbox o
    se procesara más de una imagen simultáneamente desde el mismo proceso.
    """
    from PIL import Image
    previo = Image.MAX_IMAGE_PIXELS
    Image.MAX_IMAGE_PIXELS = valor
    try:
        yield
    finally:
        Image.MAX_IMAGE_PIXELS = previo


def _validar_imagen(contenido: bytes) -> bytes:
    from PIL import Image

    try:
        with _max_image_pixels(PIXELS_MAXIMOS_IMAGEN), warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(io.BytesIO(contenido)) as img_probe:
                img_probe.verify()

            with Image.open(io.BytesIO(contenido)) as img:
                img.load()

                ancho, alto = img.size
                pixeles = ancho * alto
                if pixeles > PIXELS_MAXIMOS_IMAGEN:
                    raise ErrorValidacion(
                        f"La imagen es demasiado grande ({pixeles:,} píxeles; "
                        f"máximo {PIXELS_MAXIMOS_IMAGEN:,})."
                    )

                if getattr(img, "n_frames", 1) > 1:
                    # TIFF/WebP multipágina/animado no aporta valor aquí y amplía la
                    # superficie de parser. Solo se admiten imágenes de un único frame.
                    if img.format in {"TIFF", "WEBP"}:
                        raise ErrorValidacion("No se admiten imágenes multipágina o animadas.")

                procesada = img
                creada = None
                try:
                    if max(ancho, alto) > DIMENSION_MAXIMA_IMAGEN:
                        factor = DIMENSION_MAXIMA_IMAGEN / max(ancho, alto)
                        nuevo = (max(1, int(ancho * factor)), max(1, int(alto * factor)))
                        creada = procesada.resize(nuevo, Image.Resampling.LANCZOS)
                        procesada = creada

                    if procesada.mode not in ("RGB", "L"):
                        convertida = procesada.convert("RGB")
                        if creada is not None:
                            creada.close()
                        creada = convertida
                        procesada = convertida

                    buf = io.BytesIO()
                    procesada.save(buf, format="PNG", optimize=True)
                    normalizada = buf.getvalue()
                finally:
                    if creada is not None:
                        try:
                            creada.close()
                        except Exception:
                            pass

                if len(normalizada) > BYTES_MAXIMOS_IMAGEN_NORMALIZADA:
                    raise ErrorValidacion("La imagen normalizada es demasiado grande para procesarse.")
                return normalizada
    except ErrorValidacion:
        raise
    except Exception as e:
        raise ErrorValidacion(f"La imagen no es válida o está dañada: {e}") from None


def _limpiar_texto(texto: str) -> str:
    if len(texto) > LONGITUD_MAXIMA_TEXTO:
        logger.info("Texto truncado de %d a %d caracteres", len(texto), LONGITUD_MAXIMA_TEXTO)
        texto = texto[:LONGITUD_MAXIMA_TEXTO]

    # Mantener saltos y tabuladores, quitar controles invisibles y secuencias
    # potencialmente problemáticas para prompts/exportaciones.
    limpio = "".join(c for c in texto if c.isprintable() or c in "\n\t\r")
    limpio = re.sub(r"\n{4,}", "\n\n\n", limpio)
    return limpio.strip()


# =============================================================================
# Serialización segura del resultado del sandbox
# =============================================================================

def _documento_a_payload(doc: DocumentoProcesado) -> dict:
    return {
        "entrada": {
            "texto": doc.entrada.texto,
            "imagenes": [base64.b64encode(img).decode("ascii") for img in (doc.entrada.imagenes or [])],
            "plantilla": doc.entrada.plantilla,
            "instrucciones_tipo": doc.entrada.instrucciones_tipo,
        },
        "ruta": doc.ruta,
        "nombre_original": doc.nombre_original,
        "tipo_mime": doc.tipo_mime,
        "tamano_bytes": doc.tamano_bytes,
        "paginas": doc.paginas,
    }


def _documento_desde_payload(payload: dict) -> DocumentoProcesado:
    entrada_payload = payload.get("entrada", {})
    imagenes_b64 = entrada_payload.get("imagenes") or []
    imagenes = [base64.b64decode(img) for img in imagenes_b64]
    return DocumentoProcesado(
        entrada=Entrada(
            texto=entrada_payload.get("texto"),
            imagenes=imagenes or None,
            plantilla=entrada_payload.get("plantilla"),
            instrucciones_tipo=entrada_payload.get("instrucciones_tipo") or {},
        ),
        ruta=payload["ruta"],
        nombre_original=payload["nombre_original"],
        tipo_mime=payload["tipo_mime"],
        tamano_bytes=int(payload["tamano_bytes"]),
        paginas=payload.get("paginas"),
    )


def _procesar_impl_serializable(contenido: bytes, nombre: str, preferencia_entrada: str | None = "auto") -> dict:
    """Versión serializable para el sandbox: no devuelve objetos Python."""
    return _documento_a_payload(_procesar_impl(contenido, nombre, preferencia_entrada=preferencia_entrada))


# =============================================================================
# Función pública
# =============================================================================

def _procesar_impl(contenido: bytes, nombre: str, preferencia_entrada: str | None = "auto") -> DocumentoProcesado:
    """Implementación real del procesamiento documental.

    Esta función se ejecuta normalmente dentro de un proceso aislado.
    La preferencia de entrada viene de la interfaz y permite saltar rutas
    costosas, pero nunca evita la validación MIME real del fichero.
    """
    mime = validar(contenido, nombre)
    familia = TIPOS_ADMITIDOS[mime]
    preferencia = normalizar_preferencia_entrada(preferencia_entrada)

    if preferencia in {"pdf_texto", "pdf_escaneado"} and familia != "pdf":
        raise ErrorValidacion(
            "La opción de entrada seleccionada indica PDF, pero el fichero no es un PDF válido."
        )
    if preferencia == "imagen" and familia != "imagen":
        raise ErrorValidacion(
            "La opción de entrada seleccionada indica imagen, pero el fichero no es una imagen admitida."
        )
    if preferencia == "texto" and familia not in {"pdf", "docx", "texto"}:
        raise ErrorValidacion(
            "La opción de entrada seleccionada indica texto/PDF con capa textual, pero el fichero no contiene un formato textual admitido."
        )

    texto: str | None = None
    imagenes: list[bytes] | None = None
    ruta: Ruta
    paginas: int | None = None

    if familia == "pdf":
        texto_crudo, paginas = _extraer_texto_pdf(contenido)
        calidad = _calidad_ocr(texto_crudo, paginas)
        logger.info(
            "Calidad estimada del texto PDF: %.2f (%d chars, %d págs; preferencia=%s)",
            calidad, len(texto_crudo), paginas, preferencia,
        )

        if preferencia in {"pdf_texto", "texto"}:
            # El usuario declara que el PDF ya tiene OCR/capa textual. No se
            # renderiza ni se pasa OCR/visión. Si la capa textual no existe,
            # se devuelve un error claro para que elija "PDF escaneado".
            texto = _limpiar_texto(texto_crudo)
            if not _texto_ocr_suficiente(texto):
                raise ErrorValidacion(
                    "Ha indicado que el PDF ya tiene OCR/capa textual, pero no se ha encontrado texto suficiente. "
                    "Seleccione 'PDF escaneado/sin OCR' o 'Automático' para aplicar OCR local."
                )
            ruta = "texto"

        elif preferencia == "pdf_escaneado":
            logger.info("Preferencia de entrada: PDF escaneado; se omite la capa textual y se aplica OCR local")
            imagenes_generadas = _pdf_a_imagenes(contenido, max_paginas=PAGINAS_PDF_VISION_MAX)
            texto_ocr = _ocr_imagenes(imagenes_generadas, nombre_origen=nombre)
            if _texto_ocr_suficiente(texto_ocr):
                texto = _limpiar_texto(texto_ocr)
                imagenes = None
                ruta = "texto"
                logger.info("PDF escaneado procesado por OCR local: %d caracteres extraídos", len(texto))
            else:
                imagenes = imagenes_generadas
                ruta = "vision"
                logger.info("OCR local insuficiente en PDF escaneado; cambiando a ruta visión")
                if paginas > PAGINAS_PDF_VISION_MAX:
                    logger.info(
                        "PDF con %d páginas; se procesan las primeras %d en ruta visión",
                        paginas, PAGINAS_PDF_VISION_MAX,
                    )

        elif calidad >= 0.5:
            texto = _limpiar_texto(texto_crudo)

            if paginas > PAGINAS_PDF_HIBRIDO or not PDF_HIBRIDO_CON_IMAGENES:
                ruta = "texto"
            else:
                try:
                    imagenes = _pdf_a_imagenes(contenido, max_paginas=paginas)
                    ruta = "hibrida"
                except ErrorValidacion:
                    raise
                except Exception as e:
                    logger.warning("No se pudieron generar imágenes del PDF: %s", e)
                    ruta = "texto"
        else:
            logger.info("Texto PDF insuficiente; intentando OCR local antes de ruta visión")
            try:
                imagenes_generadas = _pdf_a_imagenes(contenido, max_paginas=PAGINAS_PDF_VISION_MAX)
                texto_ocr = _ocr_imagenes(imagenes_generadas, nombre_origen=nombre)
                if _texto_ocr_suficiente(texto_ocr):
                    texto = _limpiar_texto(texto_ocr)
                    imagenes = None
                    ruta = "texto"
                    logger.info(
                        "PDF procesado por OCR local: %d caracteres extraídos",
                        len(texto),
                    )
                else:
                    imagenes = imagenes_generadas
                    ruta = "vision"
                    logger.info("OCR local insuficiente; cambiando a ruta visión")
                    if paginas > PAGINAS_PDF_VISION_MAX:
                        logger.info(
                            "PDF con %d páginas; se procesan las primeras %d en ruta visión",
                            paginas, PAGINAS_PDF_VISION_MAX,
                        )
            except ErrorValidacion:
                raise
            except Exception as e:
                raise ErrorValidacion(
                    f"El PDF no tiene texto legible y no se pudo convertir "
                    f"a imagen/OCR para procesamiento: {e}"
                ) from None

    elif familia == "docx":
        texto = _limpiar_texto(_extraer_texto_docx(contenido))
        if not texto:
            raise ErrorValidacion("El documento DOCX no contiene texto legible.")
        ruta = "texto"

    elif familia == "texto":
        texto = _decodificar_texto(contenido)
        texto = _limpiar_texto(texto)
        if not texto:
            raise ErrorValidacion("El fichero de texto está vacío.")
        ruta = "texto"

    elif familia == "imagen":
        imagen_validada = _validar_imagen(contenido)
        texto_ocr = _ocr_imagenes([imagen_validada], nombre_origen=nombre)
        if _texto_ocr_suficiente(texto_ocr):
            texto = _limpiar_texto(texto_ocr)
            imagenes = None
            ruta = "texto"
            logger.info(
                "Imagen procesada por OCR local: %d caracteres extraídos",
                len(texto),
            )
        else:
            imagenes = [imagen_validada]
            ruta = "vision"

    else:
        raise ErrorValidacion(f"Familia de tipo no soportado: {familia}")

    entrada = Entrada(texto=texto, imagenes=imagenes)

    return DocumentoProcesado(
        entrada=entrada,
        ruta=ruta,
        nombre_original=nombre,
        tipo_mime=mime,
        tamano_bytes=len(contenido),
        paginas=paginas,
    )


def _decodificar_texto(contenido: bytes) -> str:
    for enc in ("utf-8", "cp1252", "latin-1"):
        try:
            return contenido.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ErrorValidacion("No se pudo decodificar el fichero de texto.")


def procesar(contenido: bytes, nombre: str, preferencia_entrada: str | None = "auto") -> DocumentoProcesado:
    """
    Valida y prepara el fichero para el extractor.

    Por defecto delega en un proceso hijo aislado. Esto reduce el impacto de
    bloqueos, fugas nativas o consumos anómalos en parsers de PDF/DOCX/imagen.
    Puede desactivarse con USAR_SANDBOX_PARSERS=false únicamente para depuración.
    """
    if (
        sandbox_activo()
        and os.getenv("_PLUMA_SANDBOX_CHILD") != "1"
    ):
        try:
            payload = ejecutar_en_sandbox("app.router:_procesar_impl_serializable", contenido, nombre, preferencia_entrada)
            if not isinstance(payload, dict):
                raise ErrorValidacion("El parser aislado devolvió una respuesta inválida.")
            return _documento_desde_payload(payload)
        except SandboxExecutionError as exc:
            # Preservar mensajes funcionales de validación y ocultar trazas internas.
            if exc.exception_type == "ErrorValidacion":
                raise ErrorValidacion(exc.message) from None
            logger.warning(
                "Parser aislado falló (%s): %s",
                exc.exception_type,
                exc.message,
            )
            raise ErrorValidacion(
                "No se pudo procesar el documento dentro de los límites de seguridad. "
                "Puede estar dañado, ser demasiado complejo o requerir división previa."
            ) from None

    return _procesar_impl(contenido, nombre, preferencia_entrada=preferencia_entrada)
